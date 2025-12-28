import io
from flask import Flask, render_template, request, send_file
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml import OxmlElement, ns
from pypdf import PdfReader

app = Flask(__name__)

#helper functions 

def set_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(ns.qn("w:val"), "single")
        el.set(ns.qn("w:sz"), "6")
        el.set(ns.qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def clear_p(p):
    for c in list(p._element):
        p._element.remove(c)

def write_index(cell, text=""):
    p = cell.paragraphs[0]
    clear_p(p)
    if text:
        r = p.add_run(text)
        r.bold = True
        r.font.size = Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def write_label(cell, text, indent=0.24):
    p = cell.paragraphs[0]
    clear_p(p)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(indent)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def write_value(cell, text):
    p = cell.paragraphs[0]
    clear_p(p)
    bold_tokens = [
        "REGISTERED ADDRESS:", "CORRESPONDENCE BRANCH ADDRESS:", 
        "CORRESPONDENCE ADDRESS:", "{{customer_name}}", 
        "{{client_name}}", "{{mobile}}"
    ]
    if not text:
        p.paragraph_format.left_indent = Inches(0.12)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        return

    lines = text.splitlines()
    for li, line in enumerate(lines):
        remaining = line
        while remaining:
            token = None
            for bt in bold_tokens:
                if remaining.startswith(bt):
                    token = bt; break
            if token:
                r = p.add_run(token)
                r.bold = True
                r.font.size = Pt(9)
                remaining = remaining[len(token):]
            else:
                r = p.add_run(remaining)
                r.font.size = Pt(9)
                remaining = ""
        if li != len(lines) - 1:
            p.add_run().add_break()
    p.paragraph_format.left_indent = Inches(0.12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_email(cell, email):
    p = cell.paragraphs[0]
    clear_p(p)
    part = p.part
    r_id = part.relate_to("mailto:" + email, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(ns.qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(ns.qn("w:val"), "0000FF")
    rPr.append(color)
    u = OxmlElement("w:u"); u.set(ns.qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz"); sz.set(ns.qn("w:val"), "18")
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = email
    run.append(t)
    hyperlink.append(run); p._p.append(hyperlink)
    p.paragraph_format.left_indent = Inches(0.12)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

# ------------------- CORE LOGIC -------------------

def extract_pdf_data(file_stream):
    """Parses PDF stream and extracts lines."""
    reader = PdfReader(file_stream)
    page = reader.pages[0]
    text = page.extract_text()
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    return lines

def generate_docx(pdf_lines):
    """Generates the DOCX object in memory."""
    doc = Document()
    
    # 1. Parse Data Variables from PDF lines
    HEADER_1 = pdf_lines[0]
    HEADER_2 = pdf_lines[1]
    HEADER_3 = pdf_lines[2]
    HEADER_4 = pdf_lines[3]
    HEADER_5 = pdf_lines[4]
    
    SECTION_1_HEADER = pdf_lines[5]
    ROW_1_INDEX = pdf_lines[6].split()[0]
    ROW_1_LABEL = " ".join(pdf_lines[6].split()[1:]) + "\n" + pdf_lines[7].split("{{")[0].strip()
    ROW_1_VALUE = "{{" + pdf_lines[7].split("{{")[1]
    ROW_2_HEADER = pdf_lines[8]
    ROW_3_INDEX = pdf_lines[9].split()[0]
    ROW_3_LABEL = pdf_lines[9].split()[1]
    ROW_3_VALUE = f"{pdf_lines[10]}\n{pdf_lines[11]}\n\n{pdf_lines[12]}\n{pdf_lines[13]}"
    ROW_4_LABEL, ROW_4_VALUE = "Telephone No.", "{{mobile}}"
    ROW_5_LABEL, ROW_5_VALUE = "Mobile No.", ""
    ROW_6_LABEL, ROW_6_VALUE = "Email ID", pdf_lines[16].replace("Email ID", "").strip()
    SEC_2_INDEX = pdf_lines[17].split()[0]
    SEC_2_HEADER = pdf_lines[17][2:].strip()
    SEC_2_SUB_HEADER = pdf_lines[18]
    ROW_8_LABEL, ROW_8_VALUE = "Name", "{{customer_name}}"
    ROW_9_LABEL = "Address"
    ROW_9_VALUE = (f"{pdf_lines[21]}\n{pdf_lines[22]}\n{pdf_lines[23]}\n\n"
                   f"{pdf_lines[24]}\n{pdf_lines[25]}\n{pdf_lines[26]}\n")
    DISPUTE_HEADER = pdf_lines[30]
    DISPUTE_RULE = pdf_lines[31]
    DISPUTE_TEXT = pdf_lines[32]

    # 2. Build Document
    def center(text, size, bold=True):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size)

    center(HEADER_1, 11)
    center(HEADER_2, 10)
    center(HEADER_3, 8)
    center(HEADER_4, 8, False)
    center(HEADER_5, 8, False)
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"; table.autofit = False
    
    # Section 1
    hdr = table.rows[0].cells; merged = hdr[0].merge(hdr[2])
    write_label(merged, SECTION_1_HEADER, indent=0.12); set_border(merged)
    
    r = table.add_row().cells
    write_index(r[0], ROW_1_INDEX); write_label(r[1], ROW_1_LABEL); write_value(r[2], ROW_1_VALUE)
    for c in r: set_border(c)
    
    r = table.add_row().cells
    write_index(r[0], ""); merged = r[1].merge(r[2]) # Corrected: second row index is now empty
    write_label(merged, ROW_2_HEADER); set_border(r[0]); set_border(merged)
    
    r = table.add_row().cells
    write_index(r[0], ROW_3_INDEX); write_label(r[1], ROW_3_LABEL); write_value(r[2], ROW_3_VALUE)
    for c in r: set_border(c)
    
    for l, v in [("Telephone No.", "{{mobile}}"), ("Mobile No.", ""), ("Email ID", "info@kslegal.co.in")]:
        r = table.add_row().cells
        write_index(r[0], ""); write_label(r[1], l)
        if l == "Email ID": add_email(r[2], v)
        else: write_value(r[2], v)
        for c in r: set_border(c)
        
    # Section 2
    r = table.add_row().cells
    write_index(r[0], SEC_2_INDEX); merged = r[1].merge(r[2])
    write_label(merged, SEC_2_HEADER); set_border(r[0]); set_border(merged)
    
    r = table.add_row().cells
    write_index(r[0], ""); merged = r[1].merge(r[2])
    write_label(merged, SEC_2_SUB_HEADER); set_border(r[0]); set_border(merged)
    
    rows_2 = [("Name", "{{customer_name}}"), ("Address", ROW_9_VALUE), ("Telephone No.", ""), ("Mobile No.", ""), ("Email ID", "")]
    for l, v in rows_2:
        r = table.add_row().cells
        write_index(r[0], ""); write_label(r[1], l); write_value(r[2], v)
        for c in r: set_border(c)
        
    # Dispute
    r = table.add_row().cells; merged = r[0].merge(r[2])
    write_label(merged, DISPUTE_HEADER, indent=0.12); set_border(merged)
    
    r = table.add_row().cells; merged = r[0].merge(r[2])
    p = merged.paragraphs[0]; clear_p(p)
    run = p.add_run(DISPUTE_RULE); run.bold=True; run.underline=True; run.font.size=Pt(9)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_border(merged)
    
    r = table.add_row().cells; write_index(r[0], ""); merged = r[1].merge(r[2])
    p = merged.paragraphs[0]; clear_p(p)
    run = p.add_run(DISPUTE_TEXT); run.bold=True; run.font.size=Pt(9)
    p.paragraph_format.left_indent = Inches(0.12)
    r[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_border(r[0]); set_border(merged)
    
    # Finalize
    for row in table.rows:
        try: row.cells[0].width, row.cells[1].width, row.cells[2].width = Inches(0.35), Inches(1.8), Inches(4.7)
        except: pass
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.31)
        
    return doc

# ------------------- ROUTES -------------------

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert():
    if 'file' not in request.files:
        return "No file uploaded", 400
    
    file = request.files['file']
    if file.filename == '':
        return "No selected file", 400
    
    try:
        pdf_lines = extract_pdf_data(file)
        doc = generate_docx(pdf_lines)
        output_stream = io.BytesIO()
        doc.save(output_stream)
        output_stream.seek(0)
        
        return send_file(
            output_stream,
            as_attachment=True,
            download_name="Output.docx",
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        return f"Error processing file: {str(e)}", 500

if __name__ == '__main__':
    app.run(debug=True)
